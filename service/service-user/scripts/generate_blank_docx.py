#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Generate a blank DOCX template from a filled bill-of-lading sample.

Usage:
    python3 generate_blank_docx.py <input.docx> <mappings.json> <output.docx>

The script is intentionally conservative:
- keep document/table/header/footer structure;
- replace extracted values with ${placeholder_key};
- avoid replacing static labels such as SHIPPER, CONSIGNEE, PORT OF LOADING;
- preserve run style where possible.
"""
import json
import re
import sys
from pathlib import Path

from docx import Document


KNOWN_LABELS = [
    "SHIPPER", "CONSIGNEE", "NOTIFY PARTY",
    "PRE-CARRIAGE BY", "PLACE OF RECEIPT", "OCEAN VESSEL/VOY",
    "PORT OF LOADING", "PORTC OF DISCHARGE", "PORT OF DISCHARGE",
    "PLACE OF DELIVERY", "B/L NO.", "B/L NO", "DOC. NO.", "DOC NO.",
    "BOOKING NO.", "SERVICE TYPE / MODE", "SERVICE TYPE/MODE",
    "LADEN ON BOARD", "NUMBER OF ORIGINAL B/L(S)",
    "PAYABLE AT", "PLACE AND DATE OF ISSUE",
    "FREIGHT & CHARGES", "REVENUE TONS", "RATE",
    "CONTAINER,SEAL, MARKS & NUMBER", "QUANTITY AND KIND OF PACKAGES",
    "DESCRIPTION OF GOODS", "GROSS WEIGHT (KGS)", "MEASUREMEN(M)",
    "ALSO NOTIFY PARTY (COMPLETE NAME AND ADDRESS)",
]

AMBIGUOUS_LABELS = ["PREPAID", "COLLECT", "FREIGHT PREPAID", "AS ARRANGED", "TELEX RELEASE"]

LABEL_TO_KEY = {
    "SHIPPER": "shipper",
    "CONSIGNEE": "consignee",
    "NOTIFY PARTY": "notify_party",
    "PRE-CARRIAGE BY": "pre_carriage_by",
    "PLACE OF RECEIPT": "place_of_receipt",
    "OCEAN VESSEL/VOY": "vessel_voyage",
    "PORT OF LOADING": "port_of_loading",
    "PORTC OF DISCHARGE": "port_of_discharge",
    "PORT OF DISCHARGE": "port_of_discharge",
    "PLACE OF DELIVERY": "place_of_delivery",
    "B/L NO.": "bl_no",
    "B/L NO": "bl_no",
    "DOC. NO.": "doc_no",
    "DOC NO.": "doc_no",
    "BOOKING NO.": "booking_no",
    "SERVICE TYPE / MODE": "service_type",
    "SERVICE TYPE/MODE": "service_type",
    "PAYABLE AT": "payable_at",
    "PLACE AND DATE OF ISSUE": "issue_place",
    "CONTAINER,SEAL, MARKS & NUMBER": "container_no",
    "REVENUE TONS": "revenue_tons",
    "NUMBER OF ORIGINAL B/L(S)": "original_bl_count",
    "FREIGHT & CHARGES": "freight_term",
    "DESCRIPTION OF GOODS": "goods_description",
    "GROSS WEIGHT (KGS)": "gross_weight_kgs",
    "MEASUREMEN(M)": "measurement_cbm",
    "QUANTITY AND KIND OF PACKAGES": "package_quantity",
}


def placeholder(key: str) -> str:
    return "${" + key + "}"


def normalize(text: str) -> str:
    if not text:
        return ""
    return re.sub(r"\s+", " ", text.replace("\n", " ")).strip()


def normalize_compare(text: str) -> str:
    if not text:
        return ""
    return re.sub(r"[^a-zA-Z0-9]+", "", text).lower()


def is_label(text: str, is_first_paragraph: bool = True) -> bool:
    if not text or not text.strip():
        return False
    upper = text.strip().upper()
    for label in KNOWN_LABELS:
        if upper == label or upper == label + ":" or upper == label + " :":
            return True
    if is_first_paragraph:
        for label in AMBIGUOUS_LABELS:
            if upper == label:
                return True
    is_mixed_line = ":" in text and len(text) > 40
    if (upper.startswith("FOR DELIVERY") or upper.startswith("ONWARD INLAND") or upper.startswith("ALSO NOTIFY")) and not is_mixed_line:
        return True
    if text.strip().startswith("In Witness") or text.strip().startswith("ofLading"):
        return True
    return False


def get_label_key(label: str):
    if not label:
        return None
    upper = label.strip().upper().replace(":", "").strip()
    if upper in LABEL_TO_KEY:
        return LABEL_TO_KEY[upper]
    norm = re.sub(r"[^A-Z0-9]+", "", upper)
    if norm.startswith("FORDELIVERY") or norm.startswith("ALSONOTIFY"):
        return "delivery_agent"
    return None


def surgical_replace_in_paragraph(paragraph, target: str, replacement: str) -> bool:
    runs = paragraph.runs
    if not runs or not target:
        return False

    full_text = ""
    boundaries = []
    for index, run in enumerate(runs):
        text = run.text or ""
        start = len(full_text)
        full_text += text
        boundaries.append((start, len(full_text), index))

    match_start = full_text.find(target)
    if match_start == -1:
        return False
    match_end = match_start + len(target)

    start_run_idx = None
    end_run_idx = None
    for start, end, index in boundaries:
        if start <= match_start < end:
            start_run_idx = index
        if start < match_end <= end:
            end_run_idx = index

    if start_run_idx is None or end_run_idx is None:
        return False

    start_run = runs[start_run_idx]
    start_text = start_run.text or ""
    offset_start = match_start - boundaries[start_run_idx][0]

    if start_run_idx == end_run_idx:
        prefix = start_text[:offset_start]
        suffix = start_text[offset_start + len(target):]
        start_run.text = prefix + replacement + suffix
        return True

    start_run.text = start_text[:offset_start] + replacement
    end_run = runs[end_run_idx]
    end_text = end_run.text or ""
    offset_end = match_end - boundaries[end_run_idx][0]
    end_run.text = end_text[offset_end:]
    for index in range(start_run_idx + 1, end_run_idx):
        # Keep the original run nodes so Word layout/style structure is not rebuilt.
        runs[index].text = ""
    return True


def clear_and_set_paragraph(paragraph, text: str):
    runs = paragraph.runs
    if not runs:
        if text:
            paragraph.add_run(text)
        return
    runs[0].text = text
    for index in range(1, len(runs)):
        # Preserve run/style nodes; only remove visible content.
        runs[index].text = ""


def load_mappings(mapping_path: Path):
    with mapping_path.open("r", encoding="utf-8") as file:
        raw = json.load(file)
    mappings = []
    for item in raw or []:
        original = (item.get("originalText") or item.get("original_text") or "").strip()
        key = (item.get("placeholderKey") or item.get("placeholder_key") or "").strip()
        if original and key:
            mappings.append({"originalText": original, "placeholderKey": key})
    mappings.sort(key=lambda item: -len(item["originalText"]))
    return mappings


def process_complex_delivery_cell(cell, mappings_by_key) -> bool:
    paragraphs = cell.paragraphs
    if not paragraphs or len(paragraphs) < 5:
        return False

    first_text = ""
    for paragraph in paragraphs:
        if paragraph.text and paragraph.text.strip():
            first_text = paragraph.text.strip()
            break
    if not first_text.upper().startswith("FOR DELIVERY"):
        return False

    if "delivery_agent" in mappings_by_key:
        delivery_end_idx = len(paragraphs)
        for index, paragraph in enumerate(paragraphs):
            text = paragraph.text.strip().upper() if paragraph.text else ""
            if index > 0 and (not text or text.startswith("ALSO NOTIFY")):
                delivery_end_idx = index
                break

        if delivery_end_idx > 0:
            first = paragraphs[0]
            first_text = first.text.strip()
            if ":" in first_text:
                prefix = first_text.split(":", 1)[0] + ": "
                clear_and_set_paragraph(first, prefix + placeholder("delivery_agent"))
            else:
                clear_and_set_paragraph(first, placeholder("delivery_agent"))
            for index in range(1, delivery_end_idx):
                clear_and_set_paragraph(paragraphs[index], "")

    if "carrier_agent" in mappings_by_key:
        for paragraph in paragraphs:
            text = paragraph.text.strip() if paragraph.text else ""
            upper = text.upper()
            if text.startswith("In Witness") and "ON BEHALF OF" in upper and "HAS SIGNED" in upper:
                behalf_idx = upper.find("ON BEHALF OF")
                signed_idx = upper.find("HAS SIGNED")
                if signed_idx > behalf_idx:
                    prefix_text = text[:behalf_idx + len("ON BEHALF OF")]
                    suffix_text = text[signed_idx:]
                    clear_and_set_paragraph(paragraph, prefix_text + " " + placeholder("carrier_agent") + " " + suffix_text)
                break

    return True


def process_cell(cell, mappings_by_key, mappings):
    if process_complex_delivery_cell(cell, mappings_by_key):
        process_nested_tables(cell, mappings_by_key, mappings)
        return

    first_value_idx = -1
    cell_label = ""
    found_first_non_empty = False

    for index, paragraph in enumerate(cell.paragraphs):
        text = paragraph.text
        if not text or not text.strip():
            continue
        trimmed = text.strip()
        is_first = not found_first_non_empty
        found_first_non_empty = True
        if is_label(trimmed, is_first) and first_value_idx == -1:
            cell_label = trimmed
            continue
        if first_value_idx == -1:
            first_value_idx = index

    if first_value_idx == -1:
        process_nested_tables(cell, mappings_by_key, mappings)
        return

    value_indexes = []
    value_parts = []
    for index in range(first_value_idx, len(cell.paragraphs)):
        text = cell.paragraphs[index].text
        if not text or not text.strip():
            continue
        if is_label(text.strip(), is_first_paragraph=False):
            continue
        value_indexes.append(index)
        value_parts.append(text.strip())

    if not value_parts:
        process_nested_tables(cell, mappings_by_key, mappings)
        return

    matched_placeholder = None
    label_key = get_label_key(cell_label)
    if label_key and label_key in mappings_by_key:
        matched_placeholder = placeholder(label_key)

    normalized_cell = normalize(" ".join(value_parts))
    cell_compare = normalize_compare(normalized_cell)

    if not matched_placeholder:
        for mapping in mappings:
            original = mapping["originalText"]
            original_compare = normalize_compare(original)
            if not original_compare:
                continue
            if cell_compare == original_compare:
                matched_placeholder = placeholder(mapping["placeholderKey"])
                break
            if original_compare in cell_compare and (len(original_compare) >= 3 or original_compare.isdigit()):
                replace_partial_cell_value(cell, value_indexes, original, placeholder(mapping["placeholderKey"]))

    if matched_placeholder:
        replace_whole_cell_value(cell, value_indexes, matched_placeholder)

    process_nested_tables(cell, mappings_by_key, mappings)


def replace_partial_cell_value(cell, value_indexes, original: str, replacement: str):
    first_line = original.split("\n")[0].strip()
    if not first_line:
        return
    for index in value_indexes:
        paragraph = cell.paragraphs[index]
        text = paragraph.text or ""
        if first_line in text:
            position = text.find(first_line)
            suffix = text[position + len(first_line):]
            if suffix.startswith("/"):
                clear_and_set_paragraph(paragraph, replacement)
            else:
                surgical_replace_in_paragraph(paragraph, first_line, replacement)
            return
        if len(first_line) > 10 and normalize_compare(first_line) in normalize_compare(text):
            clear_and_set_paragraph(paragraph, replacement)
            return


def replace_whole_cell_value(cell, value_indexes, replacement: str):
    first = True
    for index in value_indexes:
        paragraph = cell.paragraphs[index]
        if first:
            text = paragraph.text.strip() if paragraph.text else ""
            upper = text.upper()
            new_text = replacement
            if upper.startswith("FOR DELIVERY") or upper.startswith("ALSO NOTIFY"):
                if ":" in text:
                    new_text = text.split(":", 1)[0] + ": " + replacement
                elif " TO " in upper or upper.endswith(" TO"):
                    to_index = upper.rfind("TO")
                    new_text = text[:to_index + 2] + " " + replacement
            clear_and_set_paragraph(paragraph, new_text)
            first = False
        else:
            clear_and_set_paragraph(paragraph, "")


def process_nested_tables(cell, mappings_by_key, mappings):
    for table in cell.tables:
        for row in table.rows:
            for nested_cell in row.cells:
                process_cell(nested_cell, mappings_by_key, mappings)


def replace_in_paragraph(paragraph, mappings):
    text = paragraph.text
    if not text or not text.strip():
        return
    for mapping in mappings:
        original = mapping["originalText"]
        first_line = original.split("\n")[0].strip()
        if first_line and first_line in paragraph.text:
            surgical_replace_in_paragraph(paragraph, first_line, placeholder(mapping["placeholderKey"]))


def process_document(input_path: Path, mapping_path: Path, output_path: Path):
    document = Document(str(input_path))
    mappings = load_mappings(mapping_path)
    mappings_by_key = {mapping["placeholderKey"]: mapping for mapping in mappings}

    seen_cells = set()
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_id = id(cell._tc)
                if cell_id in seen_cells:
                    continue
                seen_cells.add(cell_id)
                process_cell(cell, mappings_by_key, mappings)

    for paragraph in document.paragraphs:
        replace_in_paragraph(paragraph, mappings)

    for section in document.sections:
        for part in [
            section.header,
            section.first_page_header,
            section.even_page_header,
            section.footer,
            section.first_page_footer,
            section.even_page_footer,
        ]:
            for paragraph in part.paragraphs:
                replace_in_paragraph(paragraph, mappings)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_path))
    print(json.dumps({"status": "success", "output": str(output_path)}, ensure_ascii=False))


def main():
    if len(sys.argv) != 4:
        raise SystemExit("Usage: generate_blank_docx.py <input.docx> <mappings.json> <output.docx>")
    process_document(Path(sys.argv[1]), Path(sys.argv[2]), Path(sys.argv[3]))


if __name__ == "__main__":
    main()
