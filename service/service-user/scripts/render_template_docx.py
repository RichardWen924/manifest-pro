#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Render a DOCX template by filling ${placeholder_key} values.

Usage:
    python3 render_template_docx.py <template.docx> <fields.json> <output.docx>

The renderer favours stable output over clever layout surgery:
- paragraphs, tables, headers and footers are visited;
- ${field_key} placeholders are replaced with Dify extracted values;
- null values become blank strings;
- multiline values are kept as Word line breaks.
"""
import json
import re
import sys
from pathlib import Path

from docx import Document


PLACEHOLDER_RE = re.compile(r"\$\{([A-Za-z0-9_]+)\}")


def load_fields(path: Path):
    with path.open("r", encoding="utf-8") as file:
        raw = json.load(file)
    fields = {}
    for key, value in (raw or {}).items():
        if value is None:
            fields[str(key)] = ""
        else:
            fields[str(key)] = str(value)
    return fields


def replace_text(text: str, fields: dict):
    def repl(match):
        key = match.group(1)
        return fields.get(key, match.group(0))

    return PLACEHOLDER_RE.sub(repl, text or "")


def surgical_replace_in_paragraph(paragraph, target: str, replacement: str) -> bool:
    runs = paragraph.runs
    if not runs or not target:
        return False

    full_text = ""
    boundaries = []
    for index, run in enumerate(runs):
        text = run.text or ""
        start = len(full_text)
        full_text += text
        boundaries.append((start, len(full_text), index))

    match_start = full_text.find(target)
    if match_start == -1:
        return False
    match_end = match_start + len(target)

    start_run_idx = None
    end_run_idx = None
    for start, end, index in boundaries:
        if start <= match_start < end:
            start_run_idx = index
        if start < match_end <= end:
            end_run_idx = index

    if start_run_idx is None or end_run_idx is None:
        return False

    start_run = runs[start_run_idx]
    start_text = start_run.text or ""
    offset_start = match_start - boundaries[start_run_idx][0]

    if start_run_idx == end_run_idx:
        prefix = start_text[:offset_start]
        suffix = start_text[offset_start + len(target):]
        set_run_text_with_breaks(start_run, prefix + replacement + suffix)
        return True

    set_run_text_with_breaks(start_run, start_text[:offset_start] + replacement)
    end_run = runs[end_run_idx]
    end_text = end_run.text or ""
    offset_end = match_end - boundaries[end_run_idx][0]
    end_run.text = end_text[offset_end:]
    for index in range(start_run_idx + 1, end_run_idx):
        runs[index].text = ""
    return True


def set_run_text_with_breaks(run, value: str):
    parts = value.splitlines()
    if not parts:
        run.text = ""
        return
    run.text = parts[0]
    for part in parts[1:]:
        run.add_break()
        run.add_text(part)


def clear_paragraph_to_text(paragraph, value: str):
    runs = paragraph.runs
    if not runs:
        set_run_text_with_breaks(paragraph.add_run(), value)
        return
    set_run_text_with_breaks(runs[0], value)
    for index in range(1, len(runs)):
        # Keep run/style nodes intact; only clear the text payload.
        runs[index].text = ""


def render_paragraph(paragraph, fields: dict):
    original = paragraph.text or ""
    if "${" not in original:
        return
    placeholders = PLACEHOLDER_RE.findall(original)
    for key in placeholders:
        if key in fields:
            surgical_replace_in_paragraph(paragraph, "${" + key + "}", fields[key])


def render_table(table, fields: dict):
    for row in table.rows:
        for cell in row.cells:
            render_block(cell, fields)


def render_block(block, fields: dict):
    for paragraph in block.paragraphs:
        render_paragraph(paragraph, fields)
    for table in block.tables:
        render_table(table, fields)


def render_document(template_path: Path, fields_path: Path, output_path: Path):
    fields = load_fields(fields_path)
    document = Document(str(template_path))
    render_block(document, fields)
    for section in document.sections:
        render_block(section.header, fields)
        render_block(section.footer, fields)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_path))


def main():
    if len(sys.argv) != 4:
        print(__doc__)
        return 2
    render_document(Path(sys.argv[1]), Path(sys.argv[2]), Path(sys.argv[3]))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
