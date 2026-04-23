#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Generate a simple editable DOCX template from Dify mappings.

Usage:
    python3 generate_placeholder_docx.py <mappings.json> <output.docx>
"""
import json
import sys
from pathlib import Path

from docx import Document


def placeholder(key: str) -> str:
    return "${" + key + "}"


def load_mappings(mapping_path: Path):
    with mapping_path.open("r", encoding="utf-8") as file:
        raw = json.load(file)
    mappings = []
    for index, item in enumerate(raw or [], start=1):
        key = (item.get("placeholderKey") or item.get("placeholder_key") or f"field_{index}").strip()
        desc = (item.get("description") or key).strip()
        original = (item.get("originalText") or item.get("original_text") or "").strip()
        mappings.append({"key": key, "description": desc, "original": original})
    return mappings


def main():
    if len(sys.argv) != 3:
        raise SystemExit("Usage: generate_placeholder_docx.py <mappings.json> <output.docx>")

    mapping_path = Path(sys.argv[1])
    output_path = Path(sys.argv[2])
    mappings = load_mappings(mapping_path)

    doc = Document()
    doc.add_heading("Bill of Lading Template", 0)
    doc.add_paragraph("Generated from extracted field mappings. Replace the layout later if a higher-fidelity source DOCX is available.")

    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Field"
    table.rows[0].cells[1].text = "Placeholder"

    for item in mappings:
        row = table.add_row().cells
        row[0].text = item["description"]
        row[1].text = placeholder(item["key"])

    doc.add_paragraph("")
    doc.add_heading("Extracted Field Reference", level=1)
    for item in mappings:
        paragraph = doc.add_paragraph()
        paragraph.add_run(item["key"]).bold = True
        paragraph.add_run(": ")
        paragraph.add_run(item["original"][:500])

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


if __name__ == "__main__":
    main()
